#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
ingest_law.py — 攝取政府法規 HTML（生產端：裸文件 → 可用語料）

支援兩種來源（依 doc_no 前綴自動選 parser）：
  - law.moj（pcode A/N…）：div 結構
  - lawweb.pcc（id 以 FL/GL 開頭）：table 結構，<br> 是排版斷行非項分隔

輸出兩種：
  ① corpus/refs_law/<法規名>.docx        ＝ 一字不漏的全文 Word（交付物）
  ② corpus/sources/<子夾>/<法規名>.md  ＝ RAG 用（frontmatter metadata + 條級本體，可進 build_catalog）

用法：python ingest_law.py <html檔> <doc_no>
"""
import re
import sys
from pathlib import Path

from docx import Document

# ── 專案根 bootstrap（搬資料夾後讓跨層 import corpus_io）──────────
_ROOT = next(p for p in Path(__file__).resolve().parents if (p / "config.example.json").exists())
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from corpus_io import SOURCES, REFS_LAW

META = {
    "A0030057": {
        "dept": "行政院公共工程委員會",
        "summary": "規範政府機關、公立學校、公營事業辦理工程、財物、勞務採購之招標、決標、履約管理、驗收、爭議處理與罰則。",
        "tags": ["政府採購", "採購", "招標", "決標", "履約管理", "驗收", "爭議處理",
                 "押標金", "保證金", "罰則", "底價", "公告金額"],
    },
    "A0030303": {
        "dept": "數位發展部",
        "summary": "規範資通安全管理法之施行細節，含公務機關與特定非公務機關之資通安全維護計畫、通報應變、稽核與演練等執行事項。",
        "tags": ["資通安全", "資安", "施行細則", "資安維護計畫", "通報應變", "稽核",
                 "演練", "公務機關", "特定非公務機關"],
    },
    "A0030304": {
        "dept": "數位發展部",
        "summary": "規範公務機關與特定非公務機關之資通安全責任等級分級基準、應辦事項與資通系統防護基準。",
        "tags": ["資通安全", "資安", "責任等級", "分級", "防護基準", "應辦事項",
                 "核心資通系統", "資安等級"],
    },
    "FL000675": {
        "dept": "行政院公共工程委員會",
        "summary": "規範機關以公開客觀評選方式委託廠商提供專業服務之評選程序與服務費用計算方式。",
        "tags": ["專業服務", "委託專業服務", "評選", "計費", "服務費用", "政府採購", "公告金額"],
    },
    "GL000077": {
        "dept": "行政院公共工程委員會",
        "summary": "規範政府採購契約應載明事項之範本，含契約文件、價金給付、履約管理、驗收、保固、權利義務與爭議處理等契約要項。",
        "tags": ["採購契約", "契約要項", "契約文件", "價金", "履約管理", "驗收",
                 "保固", "權利義務", "爭議處理", "政府採購"],
    },
    "N0030001": {
        "subdir": "政府法令",
        "dept": "勞動部",
        "summary": "規範勞動條件最低標準，含工資、工作時間、休息、休假、請假、特別休假、退休、資遣、職業災害補償、女工與童工保護等勞工權益事項。",
        "tags": ["勞動基準法", "勞基法", "工資", "工時", "休假", "特別休假", "特休",
                 "請假", "例假", "退休", "資遣", "職業災害", "加班費", "產假"],
    },
}


def strip_tags(s):
    return re.sub(r"<[^>]+>", "", s)


def minguo_to_iso(s):
    m = re.search(r"民國\s*(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日", s)
    if not m:
        return ""
    return f"{int(m.group(1)) + 1911:04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"


def parse_law_moj(html):
    name = re.search(r'id="hlLawName"[^>]*>([^<]+)</a>', html).group(1).strip()
    dm = re.search(r'id="trLNNDate">.*?<td>\s*(.*?)\s*</td>', html, re.S)
    amend = strip_tags(dm.group(1)).strip() if dm else ""
    if not minguo_to_iso(amend):  # 從未修正（無修正日期）→ fallback 公布/發布日期
        fb = re.search(r'(?:修正日期|公布日期|發布日期)：</th>\s*<td>\s*'
                       r'(民國\s*\d+\s*年\s*\d+\s*月\s*\d+\s*日)', html, re.S)
        if fb:
            amend = fb.group(1).strip()
    cm = re.search(r"法規類別：</th>\s*<td>\s*(.*?)\s*</td>", html, re.S)
    category = strip_tags(cm.group(1)).strip() if cm else ""
    start = html.find('<div class="law-reg-content">')
    body = html[start:] if start >= 0 else html
    elements = []
    for ch in re.split(r'(?=<div class="row">)', body):
        for c in re.finditer(r'<div class="h3 char-2">\s*(.*?)\s*</div>', ch, re.S):
            elements.append({"type": "chapter", "text": re.sub(r"\s+", " ", strip_tags(c.group(1))).strip()})
        nm = re.search(r'name="[\d-]+">第 (\d+(?:-\d+)?) 條', ch)
        if nm:
            seg = ch.split('<div class="col-data">', 1)
            seg = seg[1] if len(seg) > 1 else ch
            lines = [strip_tags(l).strip() for l in re.findall(r'<div class="line-[^"]*">(.*?)</div>', seg, re.S)]
            elements.append({"type": "article", "no": nm.group(1), "lines": [l for l in lines if l]})
    return {"name": name, "amend": amend, "category": category, "elements": elements}


def parse_lawweb_pcc(html):
    nm = re.search(r"法規名稱：</th>\s*<td>\s*(.*?)\s*(?:<span|</td>)", html, re.S)
    name = strip_tags(nm.group(1)).strip() if nm else ""
    am = re.search(r"修正日期：</th>\s*<td>\s*(民國\s*\d+\s*年\s*\d+\s*月\s*\d+\s*日)", html, re.S)
    if not am:
        am = re.search(r"公發布日：</th>\s*<td>\s*(民國\s*\d+\s*年\s*\d+\s*月\s*\d+\s*日)", html, re.S)
    amend = am.group(1).strip() if am else ""
    cm = re.search(r"法規體系：</th>\s*<td>\s*(.*?)\s*</td>", html, re.S)
    category = strip_tags(cm.group(1)).strip() if cm else ""
    elements = []
    for m in re.finditer(r'<td scope="row" class="th">(.*?)</td>\s*<td><div class="ClearCss">(.*?)</div></td>', html, re.S):
        th = strip_tags(m.group(1)).strip()
        raw = re.sub(r"<br\s*/?>", "\n", m.group(2)).replace("&nbsp;", " ")
        lines = [re.sub(r"[ \t]+", " ", strip_tags(l)).strip() for l in raw.split("\n")]
        lines = [l for l in lines if l]
        if not lines:
            continue
        cmo = re.search(r"第\s*(\d+(?:-\d+)?)\s*條", th)
        if cmo:  # 「第 N 條」在 th（如 FL000675）
            no, label = cmo.group(1), f"第 {cmo.group(1)} 條"
        else:  # 要項類：編號「一、」在內容首行（如 GL000077）
            hm = re.match(r"([一二三四五六七八九十百]+)、\s*(.*)", lines[0])
            if not hm:
                continue
            no, label = hm.group(1), hm.group(1) + "、"
            lines[0] = hm.group(2)
        elements.append({"type": "article", "no": no, "label": label, "lines": [l for l in lines if l]})
    return {"name": name, "amend": amend, "category": category, "elements": elements}


def make_docx(law, out_path):
    doc = Document()
    doc.add_heading(law["name"], level=0)
    doc.add_paragraph(f'修正日期：{law["amend"]}　法規體系/類別：{law["category"]}')
    for e in law["elements"]:
        if e["type"] == "chapter":
            doc.add_heading(e["text"], level=1)
        else:
            doc.add_heading(e.get("label", f'第 {e["no"]} 條'), level=2)
            for line in e["lines"]:
                doc.add_paragraph(line)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def make_rag_md(law, doc_no, src_url, out_path):
    meta = META.get(doc_no, {})
    dept = meta.get("dept") or law["category"].split("＞")[-1].strip() or law["category"]
    summary = meta.get("summary", "（待 LLM enrich）")
    tags = meta.get("tags", [])
    fm = [
        "---",
        f'name: {law["name"]}',
        f"doc_no: {doc_no}",
        f"dept: {dept}",
        f'version: "{minguo_to_iso(law["amend"])}"',
        f'effective_date: "{minguo_to_iso(law["amend"])}"',
        "expiry_date: null",
        "classification: 公開",
        "source_type: 政府法令",
        f'amend_label: {law["amend"]}',
        f"summary: {summary}",
        f'tags: [{", ".join(tags)}]',
        "status_label: 現行",
        f"source_url: {src_url}",
        "---",
        "",
        f'# {law["name"]}',
        "",
    ]
    body = []
    for e in law["elements"]:
        if e["type"] == "chapter":
            body += [f'## {e["text"]}', ""]
        else:
            body.append(f'第 {e["no"]} 條')
            body += e["lines"]
            body.append("")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(fm + body), encoding="utf-8")


if __name__ == "__main__":
    html_path, doc_no = sys.argv[1], sys.argv[2]
    html = Path(html_path).read_text(encoding="utf-8")
    if doc_no.startswith(("FL", "GL")):  # lawweb.pcc 的 id
        law = parse_lawweb_pcc(html)
        src_url = f"https://lawweb.pcc.gov.tw/LawContent.aspx?id={doc_no}"
    else:  # law.moj 的 pcode（A/N/…）
        law = parse_law_moj(html)
        src_url = f"https://law.moj.gov.tw/LawClass/LawAll.aspx?pcode={doc_no}"
    arts = sum(1 for e in law["elements"] if e["type"] == "article")
    subdir = META.get(doc_no, {}).get("subdir", "政府法規")
    docx_path = REFS_LAW / f'{law["name"]}.docx'
    md_path = SOURCES / subdir / f'{law["name"]}.md'
    make_docx(law, docx_path)
    make_rag_md(law, doc_no, src_url, md_path)
    print(f'✅ {law["name"]}（{doc_no}）：{arts} 條　dept={META.get(doc_no, {}).get("dept") or law["category"]}')
    print(f"   doc → {docx_path.relative_to(_ROOT)}")
    print(f"   md  → {md_path.relative_to(_ROOT)}")
